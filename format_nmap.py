#!/usr/bin/env python3
import os
from docx import Document
from docx.oxml.shared import OxmlElement,  qn
from docx.shared import Inches
from bs4 import BeautifulSoup
import argparse
import datetime

def shade_cell(cell, shade):
    tcPr = cell._tc.get_or_add_tcPr()
    tcVAlign = OxmlElement("w:shd")
    tcVAlign.set(qn("w:fill"), shade)
    tcPr.append(tcVAlign)

def add_if_exists(output, dictionairy, key):
    if not dictionairy:
        return

    if dictionairy.has_attr(key):
        output.append(dictionairy[key])
    else:
        output.append("")

def parse_file(xml):
    hosts = []

    for host in xml.findAll("host"):
        #Need to ask gordon about what he wants done if the host is down.
        if host.status["state"] == "down":
            continue

        addresses = []

        for address in host.findAll('address'):
            addresses.append("{} ({})".format(address["addr"], address["addrtype"]))

        hostnames = []

        for hostname in host.hostnames.findAll('hostname'):
            hostnames.append("{} ({})".format(hostname["name"], hostname["type"]))

        ports_table = []
        for port in host.ports.findAll('port'):

            port_table = []

            # Port | Port | State | Service | Reason | Product | Version | Extra info

            add_if_exists(port_table, port, "portid")
            add_if_exists(port_table, port, "protocol")
            add_if_exists(port_table, port.state, "state")
            add_if_exists(port_table, port.service, "name")
            add_if_exists(port_table, port.service, "product")
            add_if_exists(port_table, port.service, "version")

            ports_table.append(port_table)

        port_message = None

        if host.extraports:
            port_message = "The {} ports scanned but not shown above are in state: {}".format(host.extraports["count"], host.extraports["state"])

        uptime = None

        if host.uptime:
            uptime = str(datetime.timedelta(seconds=int(host.uptime["seconds"])))

        hosts.append([addresses, hostnames, ports_table, port_message, uptime])

    return hosts

def create_docx(hosts):
    doc_path = os.path.dirname(os.path.realpath(__file__)) + '/table-template.docx' if __name__ else './table-template.docx'

    document = Document(doc_path)

    document.add_heading('Nmap Results', 1)

    document.add_paragraph()

    for host in hosts:
        table = document.add_table(rows=1, cols=1)

        addresses_table = table.rows[0].cells[0].add_table(rows=1 + len(host[0]), cols=2)
        addresses_table.rows[0].cells[0].text = "Address"

        if not host[4] is None:
            addresses_table.rows[1].cells[1].text = "Uptime: {}".format(host[4])

        for idx, hostname in enumerate(host[0]):
            addresses_table.rows[1 + idx].cells[0].text = hostname

        hostname_table = table.rows[0].cells[0].add_table(rows=1 + len(host[1]), cols=1)
        hostname_table.rows[0].cells[0].text = "Hostnames"

        for idx, hostname in enumerate(host[1]):
            hostname_table.rows[1 + idx].cells[0].text = hostname

        rows = 1 + len(host[2])

        if not host[3] is None:
            rows += 1

        port_table = table.rows[0].cells[0].add_table(rows=rows, cols=6)

        port_table.rows[0].cells[0].merge(port_table.rows[0].cells[1])
        port_table.rows[0].cells[0].text = "Port"
        port_table.rows[0].cells[2].text = "State"
        port_table.rows[0].cells[3].text = "Service"
        port_table.rows[0].cells[4].text = "Product"
        port_table.rows[0].cells[5].text = "Version"


        if host[3]:
            for i in range(1,6):
                port_table.rows[-1].cells[0].merge(port_table.rows[-1].cells[i])
            port_table.rows[-1].cells[0].text = host[3]

        for i, port_list in enumerate(host[2]):
            color = "#00000"

            if port_list[2] == "open":
                color = "#EAF1DD"
            elif port_list[2] == "closed":
                color = "#F2DBDB"

            for j, port_info in enumerate(port_list):
                shade_cell(port_table.rows[1+i].cells[j], color)
                port_table.rows[1 + i].cells[j].text = port_info

        #Fix Styling
        for paragraph in table.rows[0].cells[0].paragraphs:
            paragraph.style = document.styles["NoSpacing"]

        addresses_table.style = document.styles["RedCursor"]
        hostname_table.style = document.styles["RedCursor"]
        port_table.style = document.styles["RedCursor"]

        document.add_paragraph()

    #Hide the template styles for the final document
    document.styles['RedCursor'].hidden = True
    document.styles['NoSpacing'].hidden = True

    return document

def parse_args():
    parser = argparse.ArgumentParser(description='Process Nmap XML file and produce docx table')

    parser.add_argument('infile', type=str, help='Input file (e.g. target.xml)')
    parser.add_argument('outfile', type=str, help='Output file (e.g. document.docx)')

    args = parser.parse_args()

    return (args.infile, args.outfile)

def main():
    infile_name, outfile_name = parse_args()

    file_handle = open(infile_name, 'r')

    xml = BeautifulSoup(file_handle, features="lxml")
    hosts = parse_file(xml)

    file_handle.close()

    document = create_docx(hosts)
    document.save(outfile_name)

if __name__ == '__main__':
    main()
